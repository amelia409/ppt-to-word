# 导入必要的库和模块
from docx import Document  # 用于创建和操作Word文档
from docx.shared import Pt, Inches  # 用于设置文档中的点和英寸单位
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 用于设置段落对齐方式
import io  # 用于处理二进制流
from PIL import Image  # 图像处理库
import os  # 操作系统相关功能
from .ai_writer import generate_explanation  # 导入AI写作模块，用于生成解释内容

def create_word_document(content: list, original_filename: str) -> str:
    """创建基于用户反馈改进的学术Word文档
    
    参数:
        content: 从PPT或PDF中提取的结构化内容列表
        original_filename: 原始文件名，用于生成输出文件名
        
    返回:
        生成的Word文档的完整路径
    """
    doc = Document()
    
    # 改进的样式配置
    styles = doc.styles
    font = styles['Normal'].font
    font.name = 'SimSun'  # 使用宋体或其他支持中文的字体
    font.size = Pt(12)  # 调整为适合中文阅读的字号
    
    # 专业学术封面
    doc.add_heading(f"学习文档", 0)
    doc.add_paragraph(f"基于课件：{original_filename}")
    doc.add_paragraph("使用人工智能处理，用于教学目的")
    doc.add_paragraph("内容已结构化和丰富化，便于学习\n")
    
    # 处理每个幻灯片/页面，应用关键改进
    for item in content:
        # 收集所有相关文本
        all_text = []  # 用于存储当前幻灯片的所有文本内容
        
        # 如果存在标题，则添加到文本集合中
        if item.get('title'):
            all_text.append(f"标题: {item['title']}")  # 添加标题前缀以便AI识别
        
        # 添加文本内容
        for element in item["content"]:
            # 只添加类型为文本且长度超过15个字符的内容（过滤掉太短的文本）
            if element["type"] == "text" and len(element["data"].strip()) > 15:  
                all_text.append(element["data"])
        
        # 过滤并添加图片中的文本（仅当文本有意义时）
        for img in item.get("images", []):
            # 只添加文本长度超过25个字符的图片文本（确保内容有实质性）
            if img.get("text") and len(img.get("text", "").strip()) > 25:  
                img_text = img.get("text", "").strip()
                # 过滤图片中的无关内容
                irrelevant_terms = ["slide", "页面", "背景", "模板", "设计", "点击", "这里"]
                # 检查文本中是否包含任何无关词汇
                if not any(term in img_text.lower() for term in irrelevant_terms):
                    all_text.append(f"图片内容: {img_text}")  # 添加图片文本前缀以便AI识别
        
        # 合并所有内容
        combined_content = "\n".join(all_text)  # 将所有文本用换行符连接
        
        # 只有当合并后的内容不为空时才处理
        if combined_content.strip():
            # 必须使用改进的AI生成内容
            try:
                # 调用AI写作模块生成结构化的解释内容
                enhanced_content = generate_explanation(combined_content)
                
                # 处理AI生成的具有层次结构的内容
                lines = enhanced_content.split('\n')  # 按行分割内容
                for line in lines:
                    line = line.strip()  # 去除行首尾空白
                    if not line:  # 跳过空行
                        continue
                    
                    if line.startswith('####'):
                        # 小标题（三级标题）
                        doc.add_heading(line[4:].strip(), level=3)                        
                    elif line.startswith('###'):
                        # 副标题（二级标题）
                        doc.add_heading(line[3:].strip(), level=2)
                    elif line.startswith('##'):
                        # 主标题（一级标题）
                        doc.add_heading(line[2:].strip(), level=1)

                    elif line.startswith('•') or line.startswith('- ') or line.startswith('* '):
                        # 项目符号列表（无序列表）
                        # 根据不同的项目符号类型提取文本内容
                        bullet_text = line[1:].strip() if line.startswith('•') else line[2:].strip()
                        para = doc.add_paragraph(bullet_text, style='List Bullet')  # 使用项目符号样式
                    elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
                        # 编号列表（有序列表）
                        num_text = line[3:].strip()  # 提取编号后的文本
                        para = doc.add_paragraph(num_text, style='List Number')  # 使用编号列表样式
                    elif '```' in line:
                        # 代码块标记 - 忽略标记行，但处理内容
                        continue
                    elif line.startswith('**') and line.endswith('**'):
                        # 粗体文本作为突出段落
                        para = doc.add_paragraph()
                        run = para.add_run(line[2:-2])  # 去除前后的**标记
                        run.bold = True  # 设置为粗体
                    elif len(line) > 10:  # 只处理内容实质的段落（超过10个字符）
                        # 普通段落
                        doc.add_paragraph(line)
                
            except Exception as e:
                # 如果AI处理出错，记录问题
                doc.add_heading(f"AI内容处理错误", level=1)
                doc.add_paragraph(f"错误信息: {str(e)}")  # 添加错误详情
                doc.add_paragraph("未处理的原始内容:")  # 标明以下是原始内容
                doc.add_paragraph(combined_content)  # 添加原始内容
        
        # 改进的章节分隔符（仅在有内容时添加）
        if combined_content.strip():
            doc.add_paragraph()  # 添加空白段落作为分隔
    
    # 使用改进的命名方式保存文件
    output_dir = "/app/assets/output"  # 输出目录
    os.makedirs(output_dir, exist_ok=True)  # 确保输出目录存在
    
    # 从原始文件名中提取基本名称（不含扩展名）
    base_name = os.path.splitext(original_filename)[0]
    # 构建输出文件名，使用中文前缀
    output_filename = f"学习文档_{base_name}.docx"  
    # 构建完整的输出路径
    output_path = os.path.join(output_dir, output_filename)
    
    doc.save(output_path)
    return output_path
